#import libraries required for report production
#import all libraries required to produce report
import time

#take current clock time and record as start time for file running
start_time = time.time()

import pandas as pd
import openpyxl
from openpyxl import load_workbook, workbook
from openpyxl.styles import Font 
from openpyxl.chart import BarChart, Reference 
import string
import docx
from docx.shared import Pt
import os
import time

#-------------------------------------------------------
#<<<<<     Run processing file      >>>>>
#-------------------------------------------------------

print(os.getcwd)

processing_time_start = time.time()
print("-----------------------------------------------------------------------------")
print()
print("Program has started. Pre-Processing code now running.")
print()
print("-----------------------------------------------------------------------------")

import stage_02_Objective_0_PreProcessing_004
processing_run_time = round(((time.time() - processing_time_start)/60),2)

print("-----------------------------------------------------------------------------")
print()
#print(f"Pre-Processing code run successfully in {processing_run_time} minutes. Now running DNA Machine Learning model.")
print(f"Pre-Processing code run successfully in {processing_run_time} minutes. Would you like to run the DNA Machine Learning model?")
user_decision = input(f"Enter 'y' to run the DNA ML model or 'n' to skip this section: ").lower()
print()
print("-----------------------------------------------------------------------------")

ML_model_run = False

if user_decision == 'y':
    ML_model_run = True
    print("Now starting to run DNA ML model")
    ml_time_start = time.time()
    import stage_03_Objective_1a_ML_v004
    ml_run_time = round(((time.time() - ml_time_start)/60), 2)

    print("----------------------------------------------------------")
    print()
    print(f"ML code run successfully in {ml_run_time} minutes. Now running Health Equity Audit.")
    print()
    print("----------------------------------------------------------")

else:
    print("----------------------------------------------------------")
    print()
    print(f"DNA ML model skipped. Now running Health Equity Audit.")
    print()
    print("----------------------------------------------------------")

hea_time_start = time.time()
import stage_04_Objective_1b_HEA_unmet_need_v004 #previously v003
hea_run_time = round(((time.time() - hea_time_start)/60),2)

print("--------------------------------------------------------------")
print()
print(f"HEA run successfully in {hea_run_time} minutes. Now modelling basecase carbon emissions.")
print()
print("--------------------------------------------------------------")

carbon_time_start = time.time()
import stage_05_Objective_2_carbon_emissions_v005
carbon_run_time = round(((time.time() - carbon_time_start)/60),2)

print("------------------------------------------------------------")
print()
print(f"Carbon emissions model run successfully in {carbon_run_time} minutes. Now producing maps.")
print()
print("------------------------------------------------------------")

maps_time_start = time.time()
import stage_06_Objective_2a_mapping_file_v003 #getting a NotImplementedError when trying to run file via import
maps_run_time = round(((time.time() - maps_time_start)/60),2)

print("------------------------------------------------")
print()
print(f"Maps produced in {maps_run_time} minutes. All assets created. Now creating summary report.")
print()
print("------------------------------------------------")

report_start_time = time.time()

#insert import statement for report code

#-------------------------------------------------------
#<<<<<     Produce report code starts     >>>>>
#-------------------------------------------------------

#-------------------------------------------------------
#Create functions
#-------------------------------------------------------

def text_to_arial_12(run, preferred_font_name, preferred_font_size):
    """
    A function to convert the python-docx run object's font to Arial 12 point
    This is used because the python-docx library cannot create styles, it can only
    change styles to the default or those already created in word
    """
    
    font = run.font
    
    #set style
    font.name = preferred_font_name
    font.size = Pt(preferred_font_size)
    return font

#-------------------------------------------------------

def add_appendix(doc, appendix_title, header_level, context_string, preferred_font_name, preferred_font_size):
    global appendix_count
    doc.add_page_break()
    doc.add_heading(f'Appendix {appendix_count}: {appendix_title}', header_level)
    
    paraObj = doc.add_paragraph()
    run = paraObj.add_run(context_string)
    font = text_to_arial_12(run, preferred_font_name, preferred_font_size)
    appendix_count += 1

#-------------------------------------------------------

def add_df_as_table(preferred_font_name, doc, df, style):
    # add a table to the end and create a reference variable
    # extra row is so we can add the header row

    new_table = doc.add_table(df.shape[0]+1, df.shape[1])

    # add the header rows.
    for j in range(df.shape[-1]):
        new_table.cell(0,j).text = df.columns[j]

    # add the rest of the data frame
    for i in range(df.shape[0]):
        for j in range(df.shape[-1]):
            new_table.cell(i+1,j).text = str(df.values[i,j])

    for row in new_table.rows:
        for cell in row.cells:
            paragraphs = cell.paragraphs
            for paragraph in paragraphs:
                for run in paragraph.runs:
                    font = run.font
                    font.name = preferred_font_name
                    font.size= Pt(8)

    new_table.style = style

    return new_table

#-------------------------------------------------------

#function to rename ML model outputs
def rename_ml_dna_metric_dfs(filepath):
    df_metrics = pd.read_csv(filepath)
    original_col_names = list(df_metrics.columns.values)
    original_col_names

    new_col_names = []

    for col in original_col_names:
        if "Unnamed: 0" in col:
            new_col_names.append("Metric")
        elif "Baseline" in col:
            new_col_names.append("Baseline Log Regression")
        elif "BaseLog" in col:
            new_col_names.append("k-fold Log Regression ")
        elif "Undersample" in col:
            new_col_names.append("Undersample Majority Class")
        elif "Oversample" in col:
            new_col_names.append("Oversample Minority Class")
        elif "Threshold" in col:
            new_col_names.append("Varying prediction threshold")
        elif "Smote" in col:
            new_col_names.append("Synthetic Minority Oversampling")
        elif "RegValuesNon" in col:
            new_col_names.append("Non-polynomial regularisation")
        elif "RegValuesPoly" in col:
            new_col_names.append("Polynomial regularisation")
        elif "Optuna" in col:
            new_col_names.append("Optimised Log Regression")
        else:
            print("An error has occured - column not recognised!")

    for name in range(len(new_col_names)):
        print(f"{original_col_names[name]} : {new_col_names[name]}")
    
    df_metrics.columns = new_col_names
    return df_metrics

#-------------------------------------------------------

def getFiles(dirName):

    """
    This function loops through a given directory (dirName) and returns filepaths for ALL files present, regardless of filetype, as a list. Possibly surpassed by getFilesOfSetType function now (?)
    """

    listOfFile = os.listdir(dirName)
    completeFileList = list()
    for file in listOfFile:
        completePath = os.path.join(dirName, file)
        if os.path.isdir(completePath):
            completeFileList = completeFileList + getFiles(completePath)
        else:
            completeFileList.append(completePath)

    return completeFileList

#-------------------------------------------------------

def getFilesOfSetType(dirName, type):
    
    """
    Given a directory name (dirName) and a file extension (type), this function loops over every file in the directory, and returns all filepaths of the given type in a list
    """
    
    listOfFile = os.listdir(dirName)
    completeFileList = list()
    for file in listOfFile:
        if file.endswith(type):
            completePath = os.path.join(dirName, file)
            if os.path.isdir(completePath):
                completeFileList = completeFileList + getFiles(completePath)
            else:
                completeFileList.append(completePath)
    return completeFileList

#-------------------------------------------------------

def identify_specific_filetype_by_condition_return_filepath_as_dict(list_of_conditions, list_files, desired_output_descriptor):
    """
    Function to loop through all files within a given directory location, check each file to see if it matches a given naming convention. 
    If it does, updates dictionary consisting of keys (conditions) and values (target file paths for asset of interest). 
    Returns this dictionary from function. 
    """

    dict_condition_to_desired_output_file_paths = {}

    for condition in list_of_conditions:
        for filepath in list_files:
            #if file.split('\\')[-1].split('-')[0] == output_descriptor:
            #    if file.split('\\')[-1].split('-')[1] == f'{condition}.png':
            #        print(f'Loop {counter}') #test print statement to check this is working on my machine
            #        print(file)
            #        counter +=1

            #Identify the file description text in the file name. This is the text that appears immediately after the prefix (e.g. after map001_ etc)
            file_descriptor = filepath.split('\\')[-1].split('-')[0].split('_')[1]
            #Identify the condition the file is associated with. This is the text that appears at the end of the file name, and immediately before the file extension.
            file_condition = filepath.split('\\')[-1].split('-')[-1].split('.')[0]

            #Check if both the file description and condition of interest match, if so update the dictionary with condition (key) and file path (value)
            if file_descriptor == desired_output_descriptor and file_condition == condition:
                dict_condition_to_desired_output_file_paths[condition] = filepath
    
    return(dict_condition_to_desired_output_file_paths)


#-------------------------------------------------------

#create a function to standardise the output dfs
def format_df_from_csv(
    filepath,
    group_label,
    demographic_feature_label,
    feature_name
    ):
    df_original = pd.read_csv(filepath)
    df_original.columns.values[0] = demographic_feature_label
    list_values = [feature_name for num in range(df_original.shape[0])]

    df_original.insert(0, group_label, list_values)
    return df_original

#-------------------------------------------------------

def create_condition_to_param_dict(df, desired_col_list, value_column_string, number_of_conditions, cast_as_type):
    
    """
    Function to produce a subset df from a source, and turn this into a dictionary consisting of condition as the key and the associated values of the target column name in the source df
    """
    
    subset_df = df.iloc[ : , desired_col_list]

    #convert to a dictionary
    temp_dict = subset_df.to_dict()

    #populate a dict to subsequent populate with condition (key) to prev rate multiplier (values)
    dict_condition_to_col_of_interest = {}

    for num in range(number_of_conditions):
        #extract required content from temp nested dict created above
        condition = temp_dict['condition'][num]
        col_of_interest = cast_as_type(temp_dict[value_column_string][num])

        dict_condition_to_col_of_interest[condition] = col_of_interest

    
    return dict_condition_to_col_of_interest

#-------------------------------------------------------

def change_image_size_pydoc(
    doc,
    proportion_change,
    filepath
    ):

    paraObj = doc.add_paragraph()
    run = paraObj.add_run()
    inline_shape = run.add_picture(filepath)
    
    original_width = inline_shape.width
    original_height = inline_shape.height
    #print(f"original width : {original_width}")
    #print(f"original height : {original_height}")

    new_width = int(round(original_width * proportion_change, 0))
    new_height = int(round(original_height * proportion_change, 0))
    
    inline_shape.width = new_width
    inline_shape.height = new_height

    print(f"new width : {new_width}")
    print(f"new height : {new_height}")

    return inline_shape, inline_shape.width, inline_shape.height

#-------------------------------------------------------

def add_list_fields_kept_or_dropped_to_pydoc(doc, list_of_fields, data_type, inc_or_exc):
    """
    function to add a given list of fields as a list into the report document produced using pydoc
    params:
    list_of_fields
    data type of int, str, with anything else interpreted as date time
    inc_or_exc expects a string of included or excluded. This is included in the output sentence for the report.
    """
    if data_type == 'int':
        field_type_description = 'numeric'
    elif data_type == 'str':
        field_type_description = 'text'
    else:
        field_type_description = 'date-time'

    if inc_or_exc == 'included':
        #check if any int fields replaced, if so, list in the doc, if not, state none replaced
        if len (list_of_fields) > 0:
            para = doc.add_paragraph(f"The following lists the {field_type_description} fields with missing data that were {inc_or_exc} with this missing data replaced:\n")
            for field in list_of_fields:
                run = para.add_run(f"  - {field}\n")
        else:
            para = doc.add_paragraph(f"No {field_type_description} fields with replaced missing data were {inc_or_exc}.\n")
    else:
    #check if any int fields replaced, if so, list in the doc, if not, state none replaced
        if len (list_of_fields) > 0:
            para = doc.add_paragraph(f"The following lists the {field_type_description} fields with missing data that were {inc_or_exc}:\n")
            for field in list_of_fields:
                run = para.add_run(f"  - {field}\n")
        else:
            para = doc.add_paragraph(f"No {field_type_description} fields with missing data were {inc_or_exc}.\n")
    return(para)

def go_up_one_directory():
    """
    Function to take the current working directory up one level
    """
    #starting directory
    print(os.getcwd())
    #code to go up one directory
    path_parent = os.path.dirname(os.getcwd())
    os.chdir(path_parent)
    #print revised directory to check this function has done its task as expected
    print(os.getcwd())


#-------------------------------------------------------
#set up report parameters
#-------------------------------------------------------
#identify current working director and location of the assets folder 
# containing output assets from the other .py files
cwd = os.getcwd()

#read-in user params file
filename = cwd+"\\raw_data/user_and_data_parameters/user_and_data_params.xlsx"

# read in the name of the service to which the analysis relates
service_name = pd.read_excel(filename, 'HEA_parameters', index_col=None, usecols = "C", header = 4, nrows=0)
service_name = list(service_name)[0]

#list_of_conditions = HEA_unmet_needv03.list_of_conditions

#----------------------------------------------------------------
#new code from HEA file to generate list_of_conditions variable locally in this file STARTS

#code to identify the number of conditions the user is running the tool for
filename = "raw_data/user_and_data_parameters/user_and_data_params.xlsx"
number_of_conditions_modelling_for = pd.read_excel(filename, 'HEA_condition_details', index_col=None, usecols = "D", header = 1, nrows=1)
number_of_conditions_modelling_for = int(number_of_conditions_modelling_for.iloc[0,0])

#read in the params defined by the user for each condition in scope. This will be used in subsequent steps to create a number of dictionaries required for HEA code to run.
#This approach removes a number of steps of user interaction with the code in the development environment, and enables separation of concerns in terms of sourcing ref material for the model 
# and actually running the model. 

df_condition_params_original = pd.read_excel(filename, 'HEA_condition_details', index_col=None, usecols = "B", header = 5, nrows=number_of_conditions_modelling_for)

#ID the column description ised in the params file
column_names_from_condition_df = list(df_condition_params_original.columns)

#create variables containing the text for each column name from the user params df read in above. By doing this, if there is a change to the word of the column, so long as the shape and order doesnt change
#the code will still run as we aren't hard-coding the text as a reference point.
temp_condition_col_name = column_names_from_condition_df[0]

#rename the df cols to have shorter meaningful names rather than the instructional column names from the user params file. 
df_condition_params_original.rename(columns= {
    temp_condition_col_name: "condition", 
    }, 
    inplace=True)

#create variables containing the text for the shorter column names we just assigned
revised_column_names_from_condition_df = list(df_condition_params_original.columns)

condition_col_name = revised_column_names_from_condition_df[0]

#create list_of_conditions variable
list_of_conditions = df_condition_params_original[condition_col_name].tolist()

#new code from HEA file to generate list_of_conditions variable locally in this file ends

#----------------------------------------------------------------


#get file path to top level assets directory
assets_folder = cwd+'\\'+'Assets_produced_by_code'

#create variables for each sub-directory, containing assets produced for each element of the program
#these were produced in the pre-processing file, and populated by all other .py files for the project
preprocessing_assets_path = cwd+'\\'+'Assets_produced_by_code\\01_pre_processing_assets'
hea_assets_path = cwd+'\\'+'Assets_produced_by_code\\02_HEA_assets'
dna_assets_path = cwd+'\\'+'Assets_produced_by_code\\03_DNA_ML_assets'
carbon_emissions_assets_path = cwd+'\\'+'Assets_produced_by_code\\04_Carbon_emissions_assets'

resized_processing_images_path = cwd+'\\'+'Assets_produced_by_code\\01_pre_processing_assets\\01_Resized_Images'
resized_hea_images_path = cwd+'\\'+'Assets_produced_by_code\\02_HEA_assets\\01_Resized_Images'

#image width to use when resize map images - need to test whether the outputs looks ok in the report, but also whether this holds true for short and wide LA's as opposed 
#to derbyshire which is tall and narrow. If not, need a logic to control for this.
mywidth = 500

#set parameters for document
preferred_font_name = 'Arial'
preferred_font_size = 12
appendix_count = 1 #used for incremental numbering of appendices
figure_count = 1 #used for incremental numbering of figures / charts
table_count = 1 #used for incremental numbering of tables

#dictionaries holding values that are the string descriptors in MS Word of the default table styles.
#these can be used as input to the add_df_as_table function
dict_plain_table_styles = {
    'plain_01' : 'Table Grid',
    'plain_02' : 'Table Grid Light',
    'plain_03' : 'Plain Table 1',
    'plain_04' : 'Plain Table 2',
    'plain_05' : 'Plain Table 3',
    'plain_06' : 'Plain Table 4',
    'plain_07' : 'Plain Table 5'
}

dict_grid_table_styles = {
    'grey_01': 'Grid Table 1 Light',
    'blue_01': 'Grid Table 1 Light Accent 1',
    'orange_01': 'Grid Table 1 Light Accent 2',
    'light_grey_01': 'Grid Table 1 Light Accent 3',
    'yellow_01': 'Grid Table 1 Light Accent 4',
    'teal_01': 'Grid Table 1 Light Accent 5',
    'green_01': 'Grid Table 1 Light Accent 6',
    'grey_02': 'Grid Table 2',
    'blue_02': 'Grid Table 2 Accent 1',
    'orange_02': 'Grid Table 2 Accent 2',
    'light_grey_02': 'Grid Table 2 Accent 3',
    'yellow_02': 'Grid Table 2 Accent 4',
    'teal_02': 'Grid Table 2 Accent 5',
    'green_02': 'Grid Table 2 Accent 6',
    'grey_03': 'Grid Table 3',
    'blue_03': 'Grid Table 3 Accent 1',
    'orange_03': 'Grid Table 3 Accent 2',
    'light_grey_03': 'Grid Table 3 Accent 3',
    'yellow_03': 'Grid Table 3 Accent 4',
    'teal_03': 'Grid Table 3 Accent 5',
    'green_03': 'Grid Table 3 Accent 6',
    'grey_04': 'Grid Table 4',
    'blue_04': 'Grid Table 4 Accent 1',
    'orange_04': 'Grid Table 4 Accent 2',
    'light_grey_04': 'Grid Table 4 Accent 3',
    'yellow_04': 'Grid Table 4 Accent 4',
    'teal_04': 'Grid Table 4 Accent 5',
    'green_04': 'Grid Table 4 Accent 6',
    'grey_05': 'Grid Table 5 Dark',
    'blue_05': 'Grid Table 5 Dark Accent 1',
    'orange_05': 'Grid Table 5 Dark Accent 2',
    'light_grey_05': 'Grid Table 5 Dark Accent 3',
    'yellow_05': 'Grid Table 5 Dark Accent 4',
    'teal_05': 'Grid Table 5 Dark Accent 5',
    'green_05': 'Grid Table 5 Dark Accent 6',
    'grey_06': 'Grid Table 6 Colorful',
    'blue_06': 'Grid Table 6 Colorful Accent 1',
    'orange_06': 'Grid Table 6 Colorful Accent 2',
    'light_grey_06': 'Grid Table 6 Colorful Accent 3',
    'yellow_06': 'Grid Table 6 Colorful Accent 4',
    'teal_06': 'Grid Table 6 Colorful Accent 5',
    'green_06': 'Grid Table 6 Colorful Accent 6',
    'grey_07': 'Grid Table 7 Colorful',
    'blue_07': 'Grid Table 7 Colorful Accent 1',
    'orange_07': 'Grid Table 7 Colorful Accent 2',
    'light_grey_07': 'Grid Table 7 Colorful Accent 3',
    'yellow_07': 'Grid Table 7 Colorful Accent 4',
    'teal_07': 'Grid Table 7 Colorful Accent 5',
    'green_07': 'Grid Table 7 Colorful Accent 6'

}

dict_list_table_styles = {
    'grey_01': 'List Table 1 Light',
    'blue_01': 'List Table 1 Light Accent 1',
    'orange_01': 'List Table 1 Light Accent 2',
    'light_grey_01': 'List Table 1 Light Accent 3',
    'yellow_01': 'List Table 1 Light Accent 4',
    'teal_01': 'List Table 1 Light Accent 5',
    'green_01': 'List Table 1 Light Accent 6',
    'grey_02': 'List Table 2',
    'blue_02': 'List Table 2 Accent 1',
    'orange_02': 'List Table 2 Accent 2',
    'light_grey_02': 'List Table 2 Accent 3',
    'yellow_02': 'List Table 2 Accent 4',
    'teal_02': 'List Table 2 Accent 5',
    'green_02': 'List Table 2 Accent 6',
    'grey_03': 'List Table 3',
    'blue_03': 'List Table 3 Accent 1',
    'orange_03': 'List Table 3 Accent 2',
    'light_grey_03': 'List Table 3 Accent 3',
    'yellow_03': 'List Table 3 Accent 4',
    'teal_03': 'List Table 3 Accent 5',
    'green_03': 'List Table 3 Accent 6',
    'grey_04': 'List Table 4',
    'blue_04': 'List Table 4 Accent 1',
    'orange_04': 'List Table 4 Accent 2',
    'light_grey_04': 'List Table 4 Accent 3',
    'yellow_04': 'List Table 4 Accent 4',
    'teal_04': 'List Table 4 Accent 5',
    'green_04': 'List Table 4 Accent 6',
    'grey_05': 'List Table 5 Dark',
    'blue_05': 'List Table 5 Dark Accent 1',
    'orange_05': 'List Table 5 Dark Accent 2',
    'light_grey_05': 'List Table 5 Dark Accent 3',
    'yellow_05': 'List Table 5 Dark Accent 4',
    'teal_05': 'List Table 5 Dark Accent 5',
    'green_05': 'List Table 5 Dark Accent 6',
    'grey_06': 'List Table 6 Colorful',
    'blue_06': 'List Table 6 Colorful Accent 1',
    'orange_06': 'List Table 6 Colorful Accent 2',
    'light_grey_06': 'List Table 6 Colorful Accent 3',
    'yellow_06': 'List Table 6 Colorful Accent 4',
    'teal_06': 'List Table 6 Colorful Accent 5',
    'green_06': 'List Table 6 Colorful Accent 6',
    'grey_07': 'List Table 7 Colorful',
    'blue_07': 'List Table 7 Colorful Accent 1',
    'orange_07': 'List Table 7 Colorful Accent 2',
    'light_grey_07': 'List Table 7 Colorful Accent 3',
    'yellow_07': 'List Table 7 Colorful Accent 4',
    'teal_07': 'List Table 7 Colorful Accent 5',
    'green_07': 'List Table 7 Colorful Accent 6'
}

#code to identify the string to use for the style of the tables in the final report
table_style = pd.read_excel(filename, 'Branding', index_col=None, usecols = "C", header = 0, nrows=5)

table_style_as_string = table_style.iloc[4,0]
table_style_list = table_style_as_string.split('-')

if table_style_list[0] == ".":
    chosen_style = dict_plain_table_styles[table_style_list[1]]
elif table_style_list[0] == "grid":
    chosen_style = dict_grid_table_styles[table_style_list[1]]
elif table_style_list[0] == "list":
    chosen_style = dict_list_table_styles[table_style_list[1]]
else:
    print('An error has occured with the table style import.')

#read-in user param for level_1_colour
level_1_colour = pd.read_excel(filename, 'Branding', index_col=None, usecols = "D", header = 1, nrows=0)
level_1_colour = list(level_1_colour)[0]
#read-in user param for level_2_colour
level_2_colour = pd.read_excel(filename, 'Branding', index_col=None, usecols = "D", header = 2, nrows=0)
level_2_colour = list(level_2_colour)[0]
#read-in user param for level_2_colour
level_3_colour = pd.read_excel(filename, 'Branding', index_col=None, usecols = "D", header = 3, nrows=0)
level_3_colour = list(level_3_colour)[0]
#read-in user param for level_2_colour
level_4_colour = pd.read_excel(filename, 'Branding', index_col=None, usecols = "D", header = 4, nrows=0)
level_4_colour = list(level_4_colour)[0]

#dictionaries of to convert level 1-4 colour variables containing hex colours to text colour description
dict_level_1_hex_to_colour = {
    '#003087': 'dark blue',
    '#005EB8': 'blue',
    '#0072CE': 'bright blue',
    '#41B6E6': 'light blue',
    '#00A9CE': 'aqua blue',
}

dict_level_2_hex_to_colour = {
    '#768692': 'grey',
    '#E8EDEE': 'grey',
    '#231f20': 'black',
    '#425563': 'dark grey',
    '#FFFFFF': 'white',
}

dict_level_3_hex_to_colour = {
    '#006747': 'dark green',
    '#009639': 'green',
    '#78BE20': 'light green',
    '#00A499': 'green'
}

dict_level_4_hex_to_colour = {
    '#330072': 'purple',
    '#7C2855': 'dark pink',
    '#AE2573': 'pink',
    '#8A1538': 'dark red',
    '#DA291C': 'red',
    '#ED8B00': 'orange',
    '#FFB81C': 'yellow',
    '#FAE100': 'yellow',
}


#test alternative method - read in user selection from drop down list in user params file - replaced above code. Check above no longer used, then delete above, and retain  la_of_interest_user_list_selection
la_of_interest_user_list_selection = pd.read_excel(filename, 'HEA_parameters', index_col=None, usecols = "C", header = 3, nrows=0)
la_of_interest_user_list_selection = list(la_of_interest_user_list_selection)[0]


#Identify the optimally performing machine learning logistic regression variant for later use
# TO CHECK - does this code apply regardless of what model is the optimum (believe so..)
filename_ML_top_5_higher_chance = cwd+"/Assets_produced_by_code/03_DNA_ML_assets/df001_Top5HigherChanceDna.csv"
filename_ML_top_5_lower_chance = cwd+"/Assets_produced_by_code/03_DNA_ML_assets/df002_Top5LowerChanceDna.csv"
df_ML_results_higher_chance = pd.read_csv(filename_ML_top_5_higher_chance)
df_ML_results_lower_chance = pd.read_csv(filename_ML_top_5_lower_chance)

optimal_model_name_raw = list(df_ML_results_higher_chance.columns.values)[-1]
optimal_model_name_refined = optimal_model_name_raw.split('auc_')
optimal_model_name_final = optimal_model_name_refined[-1].split('=')[0]
 
#-------------------------------------------------------
#code from HEA starts
#Read-in user parameters
#code to identify the number of conditions the user is running the tool for
filename = "raw_data/user_and_data_parameters/user_and_data_params.xlsx"
number_of_conditions_modelling_for = pd.read_excel(filename, 'HEA_condition_details', index_col=None, usecols = "D", header = 1, nrows=1)
number_of_conditions_modelling_for = int(number_of_conditions_modelling_for.iloc[0,0])

#read in the params defined by the user for each condition in scope. This will be used in subsequent steps to create a number of dictionaries required for HEA code to run.
#This approach removes a number of steps of user interaction with the code in the development environment, and enables separation of concerns in terms of sourcing ref material for the model 
# and actually running the model. 

df_condition_params_original = pd.read_excel(filename, 'HEA_condition_details', index_col=None, usecols = "B, C, D, E, F, G, H, I, J", header = 5, nrows=number_of_conditions_modelling_for)
df_condition_params_original = df_condition_params_original.fillna(1)

#ID the column description ised in the params file
column_names_from_condition_df = list(df_condition_params_original.columns)

#create variables containing the text for each column name from the user params df read in above. By doing this, if there is a change to the word of the column, so long as the shape and order doesnt change
#the code will still run as we aren't hard-coding the text as a reference point.
temp_condition_col_name = column_names_from_condition_df[0]
temp_prev_type_col_name = column_names_from_condition_df[1]
temp_numerator_col_name = column_names_from_condition_df[2]
temp_denominator_col_name = column_names_from_condition_df[3]
temp_min_age_col_name = column_names_from_condition_df[4]
temp_max_age_col_name = column_names_from_condition_df[5]
temp_gender_col_name = column_names_from_condition_df[6]

#added new cols to adjust prev multiplier rate
temp_proportion_this_need_represents = column_names_from_condition_df[7]
temp_proportion_seen_by_this_service = column_names_from_condition_df[8]

#df_condition_params_original['prevalence_multiplier'] = df_condition_params_original[temp_numerator_col_name] / df_condition_params_original[temp_denominator_col_name]
#changed prev multiplier approach to enable adjustment for specific services' context
df_condition_params_original['prevalence_multiplier'] = ((df_condition_params_original[temp_numerator_col_name] / df_condition_params_original[temp_denominator_col_name]) / df_condition_params_original[temp_proportion_this_need_represents]) * df_condition_params_original[temp_proportion_seen_by_this_service]


#rename the df cols to have shorter meaningful names rather than the instructional column names from the user params file. 
df_condition_params_original.rename(columns= {
    temp_condition_col_name: "condition", 
    temp_prev_type_col_name: "prevalence_type",
    temp_numerator_col_name: "numerator",
    temp_denominator_col_name: "denominator",
    temp_min_age_col_name: "min_age",
    temp_max_age_col_name: "max_age",
    temp_gender_col_name: "gender",
    temp_proportion_this_need_represents: "proportion_need_represents",
    temp_proportion_seen_by_this_service: "proportion_seen_by_service"
    }, 
    inplace=True)

#create variables containing the text for the shorter column names we just assigned
revised_column_names_from_condition_df = list(df_condition_params_original.columns)

condition_col_name = revised_column_names_from_condition_df[0]
prev_type_col_name = revised_column_names_from_condition_df[1]
numerator_col_name = revised_column_names_from_condition_df[2]
denominator_col_name = revised_column_names_from_condition_df[3]
min_age_col_name = revised_column_names_from_condition_df[4]
max_age_col_name = revised_column_names_from_condition_df[5]
gender_col_name = revised_column_names_from_condition_df[6]
prev_multiplier_col_name = revised_column_names_from_condition_df[9] #was 7 before

#create list_of_conditions variable
#list_of_conditions = df_condition_params_original[condition_col_name].tolist()

#create dictionary of condition to prevalence rate type
dict_condition_to_prevalence_rate = create_condition_to_param_dict(df_condition_params_original, [0, 1], prev_type_col_name, number_of_conditions_modelling_for, str)

#use above dict of condition to prevalence rate type to 
# create dictionary of condition to prevalence rate reference integer where
#1 = Age standardised rate
#2 = Crude rate
#3 = Census or no rate

dict_prev_ref_for_each_condition = {}

for condition in list_of_conditions:
    ref_num = int(dict_condition_to_prevalence_rate[condition].split(":")[0])
    dict_prev_ref_for_each_condition[condition] = ref_num

#create multiplier for prev rate multiplier
dict_prev_for_each_condition = create_condition_to_param_dict(df_condition_params_original, [0, 9], prev_multiplier_col_name, number_of_conditions_modelling_for, float) #adjusted to 9, was 7

#create dictionary of condition name to min age
dict_min_age_for_each_condition = create_condition_to_param_dict(df_condition_params_original, [0, 4], min_age_col_name, number_of_conditions_modelling_for, str)

#create dictionary of condition name to max age
dict_max_age_for_each_condition = create_condition_to_param_dict(df_condition_params_original, [0, 5], max_age_col_name, number_of_conditions_modelling_for, str)

#create dictionary of condition name to numerator
dict_numerator_for_each_condition = create_condition_to_param_dict(df_condition_params_original, [0, 2], numerator_col_name, number_of_conditions_modelling_for, int)

#create dictionary of condition name to denominator
dict_denominator_for_each_condition = create_condition_to_param_dict(df_condition_params_original, [0, 3], denominator_col_name, number_of_conditions_modelling_for, int)

#create dictionary of condition name to gender seen
dict_pop_gender_for_each_condition_text = create_condition_to_param_dict(df_condition_params_original, [0, 6], gender_col_name, number_of_conditions_modelling_for, str)
#use the above dict to create a dictionary of condition to gender integer, where
#1 = Persons
#2 = Males only
#3 = Females only

dict_pop_gender_for_each_condition = {}

for condition in list_of_conditions:
    if dict_pop_gender_for_each_condition_text[condition] == "Persons":
        dict_pop_gender_for_each_condition[condition] = 1
    elif dict_pop_gender_for_each_condition_text[condition] == "Males only":
        dict_pop_gender_for_each_condition[condition] = 2
    else:
        dict_pop_gender_for_each_condition[condition] = 3


#code from HEA ends 

#-------------------------------------------------------
#Code starts!
#-------------------------------------------------------

dict_project_stage_to_assets_file_paths = {
    'processing' : preprocessing_assets_path,
    'HEA' : hea_assets_path,
    'DNA' : dna_assets_path,
    'co2e' : carbon_emissions_assets_path 
}

list_project_stages = [
    'processing',
    'HEA',
    'DNA',
    'co2e'
    ]


#-------------------------------------------------------
#create lists for each stage of the project, containing all files within 
# the relevant project stages' assets folder (produced by preceding code)

processing_files = getFiles(preprocessing_assets_path)
hea_files = getFiles(hea_assets_path)
dna_files = getFiles(dna_assets_path)
carbon_emissions_files = getFiles(carbon_emissions_assets_path)


#-------------------------------------------------------

#Next the code will work through HEA folder, and identify all file descriptors 
# that are maps or charts (other than the IMD Decile map which will need to be 
# handled separately - this is fine though, as this is not condition specific, and 
# is instead an overall map to provide context of the LA in scope)

#using the getFilesOfSetType function, create a list consisting of all .png files within the hea assets folder
list_hea_png_file_paths = getFilesOfSetType(hea_assets_path, '.png')

#Then, use this list of png files, to create a new list containing just the file desciptors for each chart or map output from the HEA file
hea_chart_and_map_descriptors = []

for file in list_hea_png_file_paths:
    file_descriptor = file.split('\\')[-1].split('-')[0].split('_')[1]
    if file_descriptor not in hea_chart_and_map_descriptors and file_descriptor != "LsoaIMDDecile.png":
        hea_chart_and_map_descriptors.append(file_descriptor)
        

#-------------------------------------------------------
#Loop through each file description for map or chart .png files in the hea folder, in 
# the list variable created above. For each item, check whether a specific key word is 
# present in the text string (Unmet, Age, Gender, Ethnicity, or IMD). Depending on which 
# of these finite options is True, follow the relevant if-elif-else branch, and call the 
# function to retrieve the full file path for the given item for each condition in scope. 
# Return each to its own dictionary, with condition as the key and the asset file 
# path as the value.

hea_chart_and_map_descriptors_subset = []
for descriptor in hea_chart_and_map_descriptors:
    if 'MakeUpInServiceVsPopulation' in descriptor:
        hea_chart_and_map_descriptors_subset.append(descriptor)

for description in hea_chart_and_map_descriptors_subset: #this was previously a for x in hea_chart_and_map_descriptors but this failed each time for IMD as there are >1 file descriptor with IMD in it
    if 'Unmet' in description:
        dict_condition_unmet_need_map_file_paths = identify_specific_filetype_by_condition_return_filepath_as_dict(list_of_conditions, hea_files, description)
    elif 'Age' in description:
        dict_condition_Age_chart_file_paths = identify_specific_filetype_by_condition_return_filepath_as_dict(list_of_conditions, hea_files, description)
    elif 'Gender' in description:
        dict_condition_Gender_chart_file_paths = identify_specific_filetype_by_condition_return_filepath_as_dict(list_of_conditions, hea_files, description)
    elif 'Ethnicity' in description:
        dict_condition_Ethnicity_chart_file_paths = identify_specific_filetype_by_condition_return_filepath_as_dict(list_of_conditions, hea_files, description)
    elif 'IMD' in description:
        dict_condition_IMD_chart_file_paths = identify_specific_filetype_by_condition_return_filepath_as_dict(list_of_conditions, hea_files, description)
    else:
        print('File description not recognised')


for description in hea_chart_and_map_descriptors:
    if 'Unmet' in description:
        dict_condition_unmet_need_map_file_paths = identify_specific_filetype_by_condition_return_filepath_as_dict(list_of_conditions, hea_files, description)
        
#-------------------------------------------------------

#repeat similar process to above, this time to get all file paths for csv files 
# (output dataframes) from each stage of the model
#using the getFilesOfSetType function, create a list consisting of all .png files within 
# the hea assets folder

list_hea_csv_file_paths = getFilesOfSetType(hea_assets_path, '.csv')

#Then, use this list of png files, to create a new list containing just the file desciptors 
# for each chart or map output from the HEA file
hea_csv_descriptors = []

for file in list_hea_csv_file_paths:
    file_descriptor = file.split('\\')[-1].split('-')[0].split('_')[1]
    if file_descriptor not in hea_csv_descriptors and file_descriptor != 'EstimatedMetAndUnmetNeed':
        hea_csv_descriptors.append(file_descriptor)

#-------------------------------------------------------

for description in hea_csv_descriptors:
    if 'Age' in description:
        dict_condition_Age_df_file_paths = identify_specific_filetype_by_condition_return_filepath_as_dict(list_of_conditions, hea_files, description)
    elif 'Gender' in description:
        dict_condition_Gender_df_file_paths = identify_specific_filetype_by_condition_return_filepath_as_dict(list_of_conditions, hea_files, description)
    elif 'Ethnicity' in description:
        dict_condition_Ethnicity_df_file_paths = identify_specific_filetype_by_condition_return_filepath_as_dict(list_of_conditions, hea_files, description)
    elif 'IMD' in description:
        dict_condition_IMD_df_file_paths = identify_specific_filetype_by_condition_return_filepath_as_dict(list_of_conditions, hea_files, description)
    else:
        print('File description not recognised')

#-------------------------------------------------------

#need to add some logic to this to include gender results, considering gender selection by
#  user for each condition
#this code takes all the output dataframe csvs from the hea stat test between proportions, 
# reformats them so they can be concatenated, concats them vertically,
#then filters this single df to just those results that are significant, and assigns this 
# df as the value in a dictionary, where the key is the condition name.
#these can now be read into the report later on.
dict_condition_df_sig_findings = {}

for condition in list_of_conditions:
    df_age = format_df_from_csv(dict_condition_Age_df_file_paths[condition], 'Demographic', 'Category', 'Age')

    df_ethnicity = format_df_from_csv(dict_condition_Ethnicity_df_file_paths[condition], 'Demographic', 'Category', 'Ethnicity')

    df_imd = format_df_from_csv(dict_condition_IMD_df_file_paths[condition], 'Demographic', 'Category', 'Deprivation')

    if dict_pop_gender_for_each_condition_text[condition] == 'Persons':
        df_gender = format_df_from_csv(dict_condition_Gender_df_file_paths[condition], 'Demographic', 'Category', 'Gender')

        df_merged_results = pd.concat([df_age, df_ethnicity, df_imd, df_gender], axis='rows')
        mask = df_merged_results['Significant?'] == 'SIG'
        df_merged_results_sig_subset = df_merged_results[mask]
        df_merged_results_sig_subset['service_proportion'] = df_merged_results_sig_subset.service_proportion.round(3)
        df_merged_results_sig_subset['95%_CI'] = df_merged_results_sig_subset['95%_CI'].round(3)
        df_merged_results_sig_subset['population_proportion'] = df_merged_results_sig_subset['population_proportion'].round(3)
        dict_condition_df_sig_findings[condition] = df_merged_results_sig_subset
    
    df_merged_results = pd.concat([df_age, df_ethnicity, df_imd], axis='rows')
    mask = df_merged_results['Significant?'] == 'SIG'
    df_merged_results_sig_subset = df_merged_results[mask]
    df_merged_results_sig_subset['service_proportion'] = df_merged_results_sig_subset.service_proportion.round(3)
    df_merged_results_sig_subset['95%_CI'] = df_merged_results_sig_subset['95%_CI'].round(3)
    df_merged_results_sig_subset['population_proportion'] = df_merged_results_sig_subset['population_proportion'].round(3)
    dict_condition_df_sig_findings[condition] = df_merged_results_sig_subset


#-------------------------------------------------------

for condition in list_of_conditions:
    dict_condition_df_sig_findings[condition]['Sig. Finding'] = dict_condition_df_sig_findings[condition]['service_proportion'] > dict_condition_df_sig_findings[condition]['population_proportion']

    dict_condition_df_sig_findings[condition].loc[dict_condition_df_sig_findings[condition]['Sig. Finding'] == True, 'Sig. Finding'] = 'Greater than population'
    dict_condition_df_sig_findings[condition].loc[dict_condition_df_sig_findings[condition]['Sig. Finding'] == False, 'Sig. Finding'] = 'Less than population'
    dict_condition_df_sig_findings[condition].drop('Significant?', axis=1, inplace=True)
    dict_condition_df_sig_findings[condition].drop('95%_CI', axis=1, inplace=True)

#-------------------------------------------------------

#Now repeat steps above to get file paths for the ML DNA output csv files 
# consisting of the top x features associated with a dna / not dna

#To do this, repeat similar process to above, this time to get all file paths for csv files 
# (output dataframes) from each stage of the model
#using the getFilesOfSetType function, create a list consisting of all .csv files within the 
# dna assets folder
list_dna_csv_file_paths = getFilesOfSetType(dna_assets_path, '.csv')

#Then, use this list of csv files, to create a new list containing just the file desciptors for each csv output from the DNA ML file
dna_csv_descriptors = []

for file in list_dna_csv_file_paths:
    if "Top" in file:
        file_descriptor = file.split('\\')[-1].split('-')[0].split('_')[1]
        if file_descriptor not in dna_csv_descriptors:
            dna_csv_descriptors.append(file_descriptor)
        



#-------------------------------------------------------

#create and populate a dictionary with keys of Higher or Lower and values as respective 
# df for the features associated with Higher or Lower DNA chance

dict_dna_higher_lower_chance = {}
for file_name in list_dna_csv_file_paths:
    if 'Higher' in file_name.split('\\')[-1]:
        new_df = pd.read_csv(file_name)
        new_df_drop_col = new_df.iloc[:, 1:]
        dict_dna_higher_lower_chance['Higher'] = round(new_df_drop_col, 3)
    elif 'Lower' in file_name:
        new_df = pd.read_csv(file_name)
        new_df_drop_col = new_df.iloc[:, 1:]
        dict_dna_higher_lower_chance['Lower'] = round(new_df_drop_col, 3)
    else:
        #print('Could not locate the file path.')
        pass

#extract text from the now second column. This column name will always be the optimally performing model. Save the model name to a variable for later use.
optimal_dna_model_name = dict_dna_higher_lower_chance['Higher'].columns.values[1]

#-------------------------------------------------------

#Read in the ML DNA model variant performance metrics, combine into single df, 
# rename columns, transpose to fit into portrait report

#file path for ML performance metrics csv
file_name_all_models_metrics = cwd+"/Assets_produced_by_code/03_DNA_ML_assets/df003_AllModelsMetrics.csv"
#file path for ML ROC AUC metrics csv
file_name_all_models_roc_auc = cwd+"/Assets_produced_by_code/03_DNA_ML_assets/df003_AllModelsROCAUC.csv"

file_name_all_models_roc_auc_not_rounded = cwd+"/Assets_produced_by_code/03_DNA_ML_assets/df003_AllModelsROCAUCNotRounded.csv"
df_all_models_roc_auc_not_rounded = rename_ml_dna_metric_dfs(file_name_all_models_roc_auc_not_rounded)

#create df of all metrics
df_all_models_metrics = rename_ml_dna_metric_dfs(file_name_all_models_metrics)

#create df of ROC AUC
df_all_models_roc_auc = rename_ml_dna_metric_dfs(file_name_all_models_roc_auc)

#rename to give label of ROC AUC
df_all_models_roc_auc.at[0.0,'Metric'] = "ROC AUC"

#combine the 2 separate output dfs to form a single summary table
combined_df = pd.concat([df_all_models_metrics, df_all_models_roc_auc], axis=0)

#transpose this table to be better accomodated into an A4 report
final_df_of_metrics = combined_df.transpose()
final_df_of_metrics.to_csv(f"{cwd}/Assets_produced_by_code/03_DNA_ML_assets/df005_AllMetricsCombinedTransposed.csv", header=False)


#-------------------------------------------------------

#retrieve optimal model roc auc value and round to 2 dp.
roc_auc_optimal_model = float(round(df_all_models_roc_auc_not_rounded.max(axis=1),2))

if roc_auc_optimal_model >= 0.85:
    classification_strength = "high classification accuracy"
elif roc_auc_optimal_model <0.85 and roc_auc_optimal_model >0.75:
    classification_strength = "moderate classification accuracy"
else:
    classification_strength = "low classification accuracy"

#-------------------------------------------------------

dict_condition_gender_text_for_report = {}
for condition in list_of_conditions:
    if dict_pop_gender_for_each_condition_text[condition] == 'Persons':
        dict_condition_gender_text_for_report[condition] = 'all persons'
    elif dict_pop_gender_for_each_condition_text[condition] == 'Females only':
        dict_condition_gender_text_for_report[condition] = 'females only'
    else:
        dict_condition_gender_text_for_report[condition] = 'males only'

#-------------------------------------------------------
#Start of code cell to produce report now above has identified all required assets
#-------------------------------------------------------

#automate the boring stuff tutorial
#https://automatetheboringstuff.com/2e/chapter15/
#uses the demo.docx file in the working directory (copy/pasted from automate the boring stuff tutorial subdirectory)
#doc = docx.Document('demo.docx')
#len(doc.paragraphs)

execution_time = (time.time() - start_time)/60

#variable for numbering charts, lists etc
chart_counter = 1
assumption_counter = 1
parameter_counter = 1

#create new .docx file
doc = docx.Document('template.docx')

from docx.shared import Pt #new
style = doc.styles['Normal'] #new
font = style.font #new
font.name = 'Arial' #new
font.size = Pt(12) #new

#add new paragraph of text to the doc 
# note: add_paragraph and add_run accept option 2nd argument as a string of the paragraph or run object's style
run = doc.add_paragraph(f'Summary Report: Health Equity Assessment, Did Not Attend Profiling, & Patient Travel Carbon Emissions for {service_name}', 'Title')

#paraObj1 = doc.add_paragraph(
#    'This automated report summarises the findings from the decision support tool produced by Matt Eves, Anya Gopfert and Sally Brown. \n\nThe report consists of the following sections:\n'
#    )

paraObj1 = doc.add_paragraph()
run = paraObj1.add_run(f'This automated report summarises the findings from the decision support tool produced by Matthew Eves, Anya Gopfert and Sally Brown, as applied to the {service_name}. ')
run = paraObj1.add_run(f'The automated data cleaning, analysis, creation of charts, and production of this summary report was completed in {round(execution_time, 2)} minutes. \n\n')

run = paraObj1.add_run('The report consists of the following sections:\n')

run = paraObj1.add_run('1. A data quality summary for the service data used; \n')

run = paraObj1.add_run('2. HEA and a summary of modelled unmet need; \n')

run = paraObj1.add_run('3. A summary of features predictive of appointments lost through "DNA";\n')

run = paraObj1.add_run('4. Modelled patient travel carbon emissions;\n')

run = paraObj1.add_run('5. Future considerations;\n') # not sure if this section is required / fits the content of the report as blurs between focus on tool and focus on outputs for the given service?

run = paraObj1.add_run('Appendix 1. A summary of assumptions made at each stage.\n')

run = paraObj1.add_run('Appendix 2. All HEA charts identifying any areas of statistically significant variation. \n')

run = paraObj1.add_run("Appendix 3. A summary of all Machine Learning models' performance. \n")

run = paraObj1.add_run("Appendix 4. A glossary providing a brief explanation of terms used\n")

run = paraObj1.add_run("Appendix 5. Acknowledgements and references.\n")


# -----------------------------------------------------------

# -----------------------------------------------------------

#add a heading. Heading "level" can be determined by the interger param after the string
#0 = title, 1 = top header, 4 = smallest subheading
doc.add_heading('1. Data Quality Summary', 1)
paraObj2 = doc.add_paragraph(f'The chart below shows the number of missing values from each field in the source data set for this analysis. The {dict_level_3_hex_to_colour[level_3_colour]} section of the bars indicates the number of rows in the dataset for which the data was available, the {dict_level_4_hex_to_colour[level_4_colour]} section indicates the number of rows for which the data was absent. ')
run = paraObj2.add_run("Each bar represents a different demographic or appointment feature in the source data set. Where any one bar has a large proportion shaded green, this suggests a particular area of poor data quality warranting further investigation.")
paraObj2 = doc.add_paragraph("The user of the tool had the option to remove fields with missing data or retain them and 'impute' (replace) missing data with modelled values instead. Appendix 1 contains an overview of what datafields were used for each stage of the decision support tool and a record of what fields were removed or retained with imputed values.")

#add an image - e.g. chart output image from the code
#width and height arguments are optional
file_path = f'{preprocessing_assets_path}\\chart001_stacked_bar_missing_values_per_field.png'
inline_shape, inline_shape.width, inline_shape.height = change_image_size_pydoc(doc, 0.5, file_path)

# -----------------------------------------------------------
# -----------------------------------------------------------



#add a heading. Heading "level" can be determined by the interger param after the string
#0 = title, 1 = top header, 4 = smallest subheading
doc.add_heading('2. Health Equity Assessment', 1)
#insert here a line to insert the string assigned to the variable called text_condition_summary in the HEA code (this is a sentence summarising how many conditions have been modelled for, and what the condition names are)
#text_condition_summary
txt_file_location = 'Assets_produced_by_code/02_HEA_assets/text_condition_summary.txt'

with open(txt_file_location) as f:
    text_condition_summary = f.readlines()

doc.add_paragraph(text_condition_summary)

doc.add_heading('Why has an HEA taken place?', 2)
paraObj3 = doc.add_paragraph(f"The intention behind running the Health Equity Assessment was to identify any statistically significant difference between the profile of users of the {service_name} and the population of {la_of_interest_user_list_selection}. ")
run = paraObj3.add_run("Now that the areas of variance have been identified that are unlikely to be due to chance or natural variation, action planning can take place focussed on how to address the inequality/ies that have been identified. ")
run = paraObj3.add_run("A recommended process is as follows:")

doc.add_paragraph('1. Review the literature / evidence base for any factors/interventions that have been identified as associated with (or causing) higher engagement by the group(s) identified. ')
doc.add_paragraph('2. Undertake a gap analysis can be undertaken to check to what extent the service in question is operating in accordance with the evidence base for the given group(s).')
doc.add_paragraph(f'3. Where the gap analysis identifies gaps in current practice relative to the evidence base, this can form the initial evidence-based action(s) for any inequalities action plan for {service_name}.')

# -----------------------------------------------------------

#add sub-heading for this section
doc.add_heading('User-entered factors', 2)

#intro paragraph
paraObj5 = doc.add_paragraph("The following factors were entered by the user in running the model for each of the conditions in scope:\n")

#list assumptions;
for condition in list_of_conditions:

#these dictionaries contain the user params for each condition in scope. need adding to the params loop below.
#number_of_conditions_modelling_for
#df_condition_params_original
#dict_condition_to_prevalence_rate
#dict_prev_ref_for_each_condition
#dict_prev_for_each_condition
#dict_min_age_for_each_condition - DONE
#dict_max_age_for_each_condition - DONE
#dict_pop_gender_for_each_condition_text
#dict_numerator_for_each_condition
#dict_denominator_for_each_condition

#1 = Age standardised rate
#2 = Crude rate
#3 = Census or no rate

    doc.add_heading(f'{condition}:', 3)
    para = doc.add_paragraph()
    if dict_prev_ref_for_each_condition[condition] == 2: #crude rate:
        run = para.add_run(f"1. A baseline prevalence rate of {dict_numerator_for_each_condition[condition]} per {dict_denominator_for_each_condition[condition]} population.\n")
    else:
        run = para.add_run(f"1. All residents within the specified age range and gender below were included as a rate wasn't available.\n")
    run = para.add_run(f"2. A minimum and maximum age seen by the service of {dict_min_age_for_each_condition[condition]} and {dict_max_age_for_each_condition[condition]}, respectively. \n")
    run = para.add_run(f'3. The gender of patients seen for {condition} in the {service_name} is "{dict_condition_gender_text_for_report[condition].lower()}"\n')

#test section
#add sub-heading for this section
doc.add_heading('Context:', 2)
doc.add_heading('Deprivation:', 3)

#para to explain how to interpret the IMD map
doc.add_paragraph(f"The map below shows how areas of deprivation are distributed across {la_of_interest_user_list_selection}. It displays relative deprivation for each Lower Super Output Area (LSOA) using a nationally recognised measurement called the IMD Decile. A decile of 1 (dark red) indicates areas that are in the 10% most deprived LSOAs nationally, whereas a decile of 10 (dark green) indicates an area that is in the least 10% deprived LSOAs nationally.")
#doc.add_picture(f'{hea_assets_path}/map001_LsoaIMDDecile.png', width=docx.shared.Cm(10), height=docx.shared.Cm(6)) #original code

imd_map = f'{hea_assets_path}\\map001_LsoaIMDDecile.png'
inline_shape, inline_shape.width, inline_shape.height = change_image_size_pydoc(doc, 0.5, imd_map)

#end test section

doc.add_heading('Findings:', 2)
#chi square results text added 28/3
doc.add_heading('Chi-square results:', 3)
doc.add_paragraph('The following table summarises the results from chi-squared tests undertaken on each demographic in scope, for each condition, comparing the service user profile to that of the population. This provides context for the subsequent results summarised in the section that follows.')

#add the chi square results df as a table to the doc - 28/3
print(os.getcwd())
os.chdir(hea_assets_path)
df_chi_square_results = pd.read_csv('df_chi_square_results.csv')
go_up_one_directory()
go_up_one_directory()
print(os.getcwd())

add_df_as_table(preferred_font_name, doc, df_chi_square_results, chosen_style)
doc.add_paragraph()

doc.add_heading('Areas of significant variance:', 3)
doc.add_paragraph('The following table highlights the statistically significant findings of the HEA. This is found using a statistical test which compares the proportion of people with the relevant demographic in the service, to the population proportion for the Local Authority/Integrated Care System as a whole.')
doc.add_paragraph('Appendix 2 contains a series of charts, one for each condition and demographic in scope. It summarises the profile of those people who have tried to access the services compared to the population profile in the Local Authority / ICS in which the service operates. The methodology for this section is further explained in Appendix 1.')

#code for table consisting of all SIG findings - need to consider how to add gender to this in above stage (see cell further up)
for condition in list_of_conditions:
    doc.add_heading(f'{condition}:', 4)
    add_df_as_table(preferred_font_name, doc, dict_condition_df_sig_findings[condition], chosen_style)
    doc.add_paragraph()

#------------------------------------------------------------
#run a logic check to determine whether all conditions being modelled for are using a prevalence rate and not a census
# create dictionary of condition to prevalence rate reference integer where
#1 = Age standardised rate
#2 = Crude rate
#3 = Census or no rate
#all_using_rates = True
list_prev_refs_all_conditions = []
for condition in list_of_conditions:
#    if dict_prev_ref_for_each_condition[condition] == 3:
#        all_using_rates = False
    list_prev_refs_all_conditions.append(dict_prev_ref_for_each_condition[condition])

list_unique_prev_refs_all_conditions = list(set(list_prev_refs_all_conditions))

#Bool to check all conditions using crude rate, or crude rate plus one of the other 2 options
all_using_crude_rate = False
all_using_crude_or_age_standardised_rates = False
all_using_crude_or_census = False
#Bool to check all conditions using age standardised, or, age standardised plus census (age standard plus crude already handled above)
all_using_age_standardised = False
all_using_age_standardised_or_census = False
#Bool to check all conditions using census approach only
all_using_census = False
#
all_using_combo_crude_age_standardised_and_census = False

if len(list_of_conditions) > 1:
    plural = "s"
    single_plural = ""
else:
    plural = ""
    single_plural = "s"

#only age standardised
if 1 in list_unique_prev_refs_all_conditions and len(list_unique_prev_refs_all_conditions) == 1:
    model_rate_summary = 'age standardised prevalence rate'
    all_using_age_standardised = True

#only crude rates
elif 2 in list_unique_prev_refs_all_conditions and len(list_unique_prev_refs_all_conditions) == 1:
    model_rate_summary = 'crude prevalence rate'
    all_using_crude_rate = True

#only census
elif 3 in list_unique_prev_refs_all_conditions and len(list_unique_prev_refs_all_conditions) == 1:
    model_rate_summary = f'Prevalence rates have not been used. Instead the condition(s) in scope have been modelled using the total target population in the age range as applicable to the patient gender(s) seen by the service. '
    all_using_census = True

#combo of crude and age standardised rates
elif 1 in list_unique_prev_refs_all_conditions and 2 in list_unique_prev_refs_all_conditions and len(list_unique_prev_refs_all_conditions) == 2:
    model_rate_summary = 'combination of age standardised and crude prevalence rates'
    all_using_crude_or_age_standardised_rates = True

#combo of crude rate and census
elif 2 in list_unique_prev_refs_all_conditions and 3 in list_unique_prev_refs_all_conditions  and len(list_unique_prev_refs_all_conditions) == 2:
    model_rate_summary = 'combination of crude prevalence rate(s) and the use of a census (total target population)'
    all_using_crude_or_census = True

#combo of age standardised and census
elif 1 in list_unique_prev_refs_all_conditions and 3 in list_unique_prev_refs_all_conditions  and len(list_unique_prev_refs_all_conditions) == 2:
    model_rate_summary = 'combination of age standardised rate(s) and the use of a census (total target population)'
    all_using_age_standardised_or_census = True

#combo of age standardised, crude rate, and census
elif 1 in list_unique_prev_refs_all_conditions and 2 in list_unique_prev_refs_all_conditions and 3 in list_unique_prev_refs_all_conditions:
    model_rate_summary = 'combination of age standardised rate(s), crude rate(s) and the use of a census (total target population)'
    all_using_combo_crude_age_standardised_and_census = True

else:
    pass




#add conditional logic to format the doc content appropriately, to reflect whether the condition(s) modelled are all using rates or not
if all_using_age_standardised == True:
    doc.add_heading('Modelled unmet need:', 3)
    para = doc.add_paragraph()
    run = para.add_run("The tool was used to estimate the level of unmet need in the population for each condition in scope, based on the methodology and assumptions outlined in Appendix 2. ")
    if len(list_of_conditions) >1:
        run = para.add_run(f"To do this, a {model_rate_summary} has been used for each of the {str(len(list_of_conditions))} conditions modelled. ")
    else:
        run = para.add_run(f"To do this, an {model_rate_summary} has been used.")
    run = para.add_run(f"The following map{plural} indicate{single_plural} the level (and location) of modelled unmet need. ")
    run = para.add_run(f"The map{plural} indicate{single_plural} areas where prevalence would suggest there are the highest numbers of residents likely to have the condition of interest but not seen by the service, in dark red, and the areas likely to have the lowest number of residents with the condition of interest but not accessed the service, in very pale red (almost white). ")

elif all_using_crude_rate == True:
    doc.add_heading('Modelled unmet need:', 3)
    para = doc.add_paragraph()
    run = para.add_run("The tool was used to estimate the level of unmet need in the population for each condition in scope, based on the methodology and assumptions outlined in Appendix 2. ")
    if len(list_of_conditions) >1:
        run = para.add_run(f"To do this, a {model_rate_summary} has been used for each of the {str(len(list_of_conditions))} conditions modelled. ")
    else:
        run = para.add_run(f"To do this, a crude prevalence rate has been used.")
    run = para.add_run(f"The following map{plural} indicate{single_plural} the level (and location) of modelled unmet need. ")
    run = para.add_run(f"The map{plural} indicate{single_plural} areas where prevalence would suggest there are the highest numbers of residents likely to have the condition of interest but not seen by the service, in dark red, and the areas likely to have the lowest number of residents with the condition of interest but not accessed the service, in very pale red (almost white). ")

elif all_using_census == True:
    doc.add_heading('Modelled population not utilising the service:', 3)
    para = doc.add_paragraph()
    run = para.add_run("The tool was used to estimate the number of people in the local population who have not utilised the service, for each condition in scope, based on the methodology and assumptions outlined in Appendix 2. ")
    run = para.add_run(f"{model_rate_summary}")
    run = para.add_run(f"The following map{plural} indicate{single_plural} the size of the population eligible to access the service that hasn't accessed it in the timeframe in scope for the analysis. For clarity, this approach was used in the absence of a prevalence rate. Re-running the analysis with the use of a prevalence rate would be recommended, if possible. ")
    run = para.add_run(f"The map{plural} indicate{single_plural} areas where there are the highest numbers of residents eligible to access the service but haven't done so, in dark red, and the areas with the least number of residents eligible to access the service that haven't done so in very pale red (almost white). ")

elif all_using_crude_or_age_standardised_rates == True:
    doc.add_heading('Modelled unmet need:', 3)
    para = doc.add_paragraph()
    run = para.add_run("The tool was used to estimate the level of unmet need in the population for each condition in scope, based on the methodology and assumptions outlined in Appendix 2. ")
    if len(list_of_conditions) >1:
        run = para.add_run(f"To do this, a {model_rate_summary} has been used for the {str(len(list_of_conditions))} conditions modelled. ")
    else:
        run = para.add_run(f"To do this, a {model_rate_summary} has been used.")
    run = para.add_run(f"The following map{plural} indicate{single_plural} the level (and location) of modelled unmet need. ")
    run = para.add_run(f"The map{plural} indicate{single_plural} areas where prevalence would suggest there are the highest numbers of residents likely to have the condition of interest but not seen by the service, in dark red, and the areas likely to have the lowest number of residents with the condition of interest but not accessed the service, in very pale red (almost white). ")
    
elif all_using_crude_or_census == True:
    doc.add_heading('Modelled unmet need:', 3)
    para = doc.add_paragraph()
    run = para.add_run("The tool was used to estimate the level of unmet need in the population for each condition in scope, based on the methodology and assumptions outlined in Appendix 2. ")
    if len(list_of_conditions) >1:
        run = para.add_run(f"To do this, a {model_rate_summary} has been used for the {str(len(list_of_conditions))} conditions modelled. ")
    else:
        run = para.add_run(f"To do this, a {model_rate_summary} has been used.")

    run = para.add_run(f"Where prevalence has been used, the following map{plural} indicate{single_plural} the level (and location) of modelled unmet need, with the number of residents likely to have the condition of interest but not seen by the service, in dark red, and the areas likely to have the lowest number of residents with the condition of interest but not accessed the service, in very pale red (almost white). ")
    run = para.add_run(f"Where a census has been used, the map{plural} indicate{single_plural} the size of the population eligible to access the service who have not accessed it in the timeframe in scope for the analysis. For clarity, a census rate was used in the absence of a prevalence rate. ")
    run = para.add_run(f"In this situation, the map{plural} indicate{single_plural} areas where there are the highest numbers of residents eligible to access the service but haven't done so, in dark red, and the areas with the least number of residents eligible to access the service that haven't done so in very pale red (almost white).")
    run = para.add_run(f"Re-running the analysis with the use of a prevalence rate would be recommended, if possible. ")

elif all_using_age_standardised_or_census == True:
    doc.add_heading('Modelled unmet need:', 3)
    para = doc.add_paragraph()
    run = para.add_run("The tool was used to estimate the level of unmet need in the population for each condition in scope, based on the methodology and assumptions outlined in Appendix 2. ")
    if len(list_of_conditions) >1:
        run = para.add_run(f"To do this, a {model_rate_summary} has been used for the {str(len(list_of_conditions))} conditions modelled. ")
    else:
        run = para.add_run(f"To do this, a {model_rate_summary} has been used.")

    run = para.add_run(f"Where prevalence has been used, the following map{plural} indicate{single_plural} the level (and location) of modelled unmet need, with the number of residents likely to have the condition of interest but not seen by the service, in dark red, and the areas likely to have the lowest number of residents with the condition of interest but not accessed the service, in very pale red (almost white). ")
    run = para.add_run(f"Where a census has been used, the map{plural} indicate{single_plural} the size of the population eligible to access the service who have not accessed it in the timeframe in scope for the analysis. For clarity, a census rate was used in the absence of a prevalence rate. ")
    run = para.add_run(f"In this situation, the map{plural} indicate{single_plural} areas where there are the highest numbers of residents eligible to access the service but haven't done so, in dark red, and the areas with the least number of residents eligible to access the service that haven't done so in very pale red (almost white).")
    run = para.add_run(f"Re-running the analysis with the use of a prevalence rate would be recommended, if possible. ")

elif all_using_combo_crude_age_standardised_and_census == True:
    doc.add_heading('Modelled unmet need:', 3)
    para = doc.add_paragraph()
    run = para.add_run("The tool was used to estimate the level of unmet need in the population for each condition in scope, based on the methodology and assumptions outlined in Appendix 2. ")
    if len(list_of_conditions) >1:
        run = para.add_run(f"To do this, a {model_rate_summary} has been used for the {str(len(list_of_conditions))} conditions modelled. ")
    else:
        run = para.add_run(f"To do this, a {model_rate_summary} has been used.")

    run = para.add_run(f"Where prevalence has been used, the following map{plural} indicate{single_plural} the level (and location) of modelled unmet need, with the number of residents likely to have the condition of interest but not seen by the service, in dark red, and the areas likely to have the lowest number of residents with the condition of interest but not accessed the service, in very pale red (almost white). ")
    run = para.add_run(f"Where a census has been used, the map{plural} indicate{single_plural} the size of the population eligible to access the service who have not accessed it in the timeframe in scope for the analysis. For clarity, a census rate was used in the absence of a prevalence rate. ")
    run = para.add_run(f"In this situation, the map{plural} indicate{single_plural} areas where there are the highest numbers of residents eligible to access the service but haven't done so, in dark red, and the areas with the least number of residents eligible to access the service that haven't done so in very pale red (almost white).")
    run = para.add_run(f"Re-running the analysis with the use of a prevalence rate would be recommended, if possible. ")

else:
    pass

    '''
    doc.add_heading('Modelled need:', 3)
    
    doc.add_paragraph("The tool was used to estimate the level of unmet need in the population for each condition in scope, based on the methodology and assumptions outlined in Appendix 2. ")
    para = doc.add_paragraph()
    run = para.add_run("Unfortunately, prevalence rates were not available for all conditions included in the scope of this work.")
    run = para.add_run("The following maps indicate where a prevalence rate has been used (if at all). In these circumstances, the maps the level (and location) of modelled unmet .")
    doc.add_paragraph("The following map(s) indicate the level (and location) of modelled unmet need by showing the areas where the data indicates there are the highest number of residents likely to have the condition of interest in dark red, and the areas with the lowest number of residents likely to have the condition of interest in very pale red (almost white).")
    '''

for condition in list_of_conditions:
    doc.add_heading(f'{condition}:', 3)
    inline_shape, inline_shape.width, inline_shape.height = change_image_size_pydoc(doc, 0.5, dict_condition_unmet_need_map_file_paths[condition])

#-----------------------------------------------------------

#add a heading. Heading "level" can be determined by the interger param after the string
#0 = title, 1 = top header, 4 = smallest subheading
doc.add_heading('3. Features predictive of a DNA', 1)
'''
doc.add_paragraph()
run = paraObj2.add_run("The tool in its current form allows the user to either remove a data field altogether, or, impute (replace) missing data with an alternative. For text fields this replacement method involves simply inserting the word 'missing' for any missing data item.")
run = paraObj2.add_run("For number fields (such as age) the current approach involves replace the missing data with the median (middle) value of that data field.")
'''

if user_decision == 'y':
    doc.add_paragraph("The chart below displays the sample size used and expected predictive accuracy associated with this sample size (along with other possible sample sizes). The red cross highlights the sample size used for the model and the anticipated 'out of the box' accuracy that may be expected with the initial base logistic regression model for that sample size. It is worth noting accuracy will increase with increased sample size, and the choice to use a sample size less than the total is usually due to limitations such as insufficient computer power to run the model (though this is not an anticipated issue for the model created). ")
    file_path = f'{dna_assets_path}\\chart001_FigSelectedSampleSizeAccuracy.png'
    doc.add_picture(file_path, width=docx.shared.Cm(10), height=docx.shared.Cm(6))

    doc.add_paragraph(f"The model used for this component was a logistic regresssion. 9 variants of this logistic regression were run with the best performing being {optimal_model_name_final} with a ROC Area Under the Curve value of {roc_auc_optimal_model} and {classification_strength}.")
    doc.add_paragraph("The performance metrics for this top performing variant are summarised in the table below and the metrics for all models (including the other 8 alternative variants) are in Appendix 3.")

    #For loop to add the tables to the doc outlining the top 5 features associated with higher or lower DNA chance, respectively.
    for direction in ['Higher', 'Lower']:
        doc.add_heading(f'Features associated with a {direction.lower()} chance of non-attendance:', 3)
        #para = doc.add_paragraph()
        add_df_as_table(preferred_font_name, doc, dict_dna_higher_lower_chance[direction], chosen_style)
        run = para.add_run()

else:
    doc.add_paragraph("There are no outputs to report in this section as the machine learning models were not run on the dataset.")

#-----------------------------------------------------------

doc.add_heading('4. Modelled carbon emissions: ', 1)
doc.add_paragraph("The tool has been used to model carbon emissions due to patient travel to and from face-to-face appointments using the standardised measure of Carbon Dioxide Equivalents (CO2e). Using these units allows comparison between different carbon producting processes.")
#ADD Measure for the estimated carbon emissions for the service.

#text_condition_summary
txt_file_location = 'Assets_produced_by_code/04_Carbon_emissions_assets/TotalEmissionsBaseline.txt'

with open(txt_file_location) as f:
    carbon_emissions_summary = f.readlines()

doc.add_paragraph(carbon_emissions_summary)


#-----------------------------------------
#section 5 future considerations

doc.add_heading('5. Future considerations: ', 1)
doc.add_paragraph("The creators of the decision support tool have identified several future developments that they believe would add value to the tool and support decision-making. These are summarised below and engagement on these is welcomed.")

# modelling new clinic location impact on emissions and unmet need
paraObj2 = doc.add_paragraph()
run = paraObj2.add_run("Incorporating modelled impact of a new clinic location ").bold=True
run = paraObj2.add_run("This could involve modelling the potential impact of a new clinic location on both unmet need and modelled emissions. This was the original aim of the work behind the creation of the decision support tool for which the foundations are now in place.")

# expanding classification algorithms for DNA ML
paraObj2 = doc.add_paragraph()
run = paraObj2.add_run("Expanding the DNA Machine Learning model ").bold=True
run = paraObj2.add_run("Currently the DNA Machine Learning model uses variations of a logistic regression model. This could (and likely, should) be expanded to include additional classification algorithms such as random forests or the use of a neural net model to determine whether these yield a better prediction.")

# Incorporating census 2021 data / improving scope of the demographics included
paraObj2 = doc.add_paragraph()
run = paraObj2.add_run("Expanding the demographics in scope ").bold=True
run = paraObj2.add_run("Currently the HEA statistical test between proportions includes a subset of the protected characteristics only due to limitations in available data at the population level. Additionally, for gender, which is included, the available data is limited and not reflective of the gender identities in society. Depending on the content and outputs from the latest census, this may be able to be adjusted for / improved upon.")

#  reducing user interaction
paraObj2 = doc.add_paragraph()
run = paraObj2.add_run("Reducing user-to-code interaction ").bold=True
run = paraObj2.add_run("While every effort has been made to minimise the level of user interaction with the code via the command line, this has not been totally removed. Further work could reduce the level of interaction with the aim of full automation and / or incoporating a Graphical User Interface (GUI), similar to how one interacts with other programs on the computer, like Word, Excel etc. to make using the tool easier for those unfamiliar with code/coding. It is anticipated this would aid scaling the use of the tool.")

# adding age standardised rates
paraObj2 = doc.add_paragraph()
run = paraObj2.add_run("Incorporating age-standardised rates ").bold=True
run = paraObj2.add_run("Currently the tool can accept a crude rate or, where this is not available, the tool can be instructed to use a census approach taking the whole population of a given gender and between a min and max age. By incorporating the option to use age-standardised rates this could scale the potential use cases.")

# data limitations
paraObj2 = doc.add_paragraph()
run = paraObj2.add_run("Considering solutions to data availability limitations ").bold=True
run = paraObj2.add_run("For certain services, where out of area patient flow is a feature, the service will likely not have the data for their residents that travelled to an out-of-area provider to be seen. This will likely be a current data/knowledge gap for the service, and this tool can't solve that. However, where services submit their data to National bodies (such as UKHSA or NHS Digtal, as at 2022) there is the potential to explore the use of this tool by those bodies.")


#ADD Measure for the estimated carbon emissions for the service.


#-----------------------------------------
#Appendix 1 - summary of method and assumptions for each stage
add_appendix(doc, "Summary of methods used", 1, 'This appendix includes a summary overview of the method and assumptions underpinning each component of the tool.', preferred_font_name, preferred_font_size)

doc.add_heading("Summary of fields used in each stage", 2)
doc.add_paragraph("The pre-processing stage of the tool 'chunks up' the raw data into 3 separate data sets to start with, 1 for each stage of the program, before each is subsequently further processed to get to a processed data file for each stage. The table below summarises the fields in scope at the initial stage for each component.\n")
add_df_as_table(preferred_font_name, doc, stage_02_Objective_0_PreProcessing_004.df_summary_of_fields_used_in_each_section, chosen_style)

para = doc.add_paragraph("\nBroadly speaking, each final processed data file created:\n")
run = para.add_run("1. Includes fields that originally had no missing data;\n")
run = para.add_run("2. Excludes fields that had missing data, and for which the user decided to drop from the data set and not impute (create) data for;\n")
run = para.add_run("3. Includes fields that originally had missing data, and for which the user imputed (created) data. \n")
run = para.add_run("4. Where missing data was imputed, 2 strategies are used. For numeric data, the median of that field is used. For text data, the word 'missing' is used.\n")

#sub-heading HEA
doc.add_heading('01. HEA', 3)

#methods section
doc.add_heading("Data set description:", 4)
para = doc.add_paragraph("The final processed dataset used for the HEA consists of: \n")
run = para.add_run("1. Unique patients (most recent appointment)\n")
run = para.add_run("2. All appointment statuses and not just those patients that attended their appointment. The HEA can therefore be viewed as a summary of 'expressed need'.\n")
run = para.add_run("3. 'In area' attendances only. This is because the HEA compares to the chosen local population profile.\n")

#list imputed fields
doc.add_heading("Approach to handling missing data:", 5)
doc.add_heading("Fields imputed (missing data replaced):", 5)
para = add_list_fields_kept_or_dropped_to_pydoc(doc, stage_02_Objective_0_PreProcessing_004.hea_for_report_int_float_fields_keep, 'int', 'included')
para = add_list_fields_kept_or_dropped_to_pydoc(doc, stage_02_Objective_0_PreProcessing_004.hea_for_report_str_fields_keep, 'str', 'included')
para = add_list_fields_kept_or_dropped_to_pydoc(doc, stage_02_Objective_0_PreProcessing_004.hea_for_report_datetime_fields_keep, 'date', 'included')

#list dropped fields with missing data
doc.add_heading("Fields with missing data that were removed from the analysis:", 4)
para = add_list_fields_kept_or_dropped_to_pydoc(doc, stage_02_Objective_0_PreProcessing_004.hea_for_report_int_float_fields_drop, 'int', 'excluded')
para = add_list_fields_kept_or_dropped_to_pydoc(doc, stage_02_Objective_0_PreProcessing_004.hea_for_report_str_fields_drop, 'str', 'excluded')
para = add_list_fields_kept_or_dropped_to_pydoc(doc, stage_02_Objective_0_PreProcessing_004.hea_for_report_datetime_fields_drop, 'date', 'excluded')

#sub-heading DNA ML section
doc.add_heading('02. DNA', 3)

#methods section
doc.add_heading("Data set description:", 4) #zxcv continue from here to build out methods text 
para = doc.add_paragraph("The final processed dataset used for the DNA Machine Learning model consists of: \n")
run = para.add_run("1. Unique appointments (not patients)\n")
run = para.add_run("2. Demographic fields (but not LSOA as currently this causes the model not to run as expected).\n")
run = para.add_run("3. 'In area' attendances only.\n")
run = para.add_run("4. Imputed (replaced) missing data wherever the user requested this (as indicated in the following section).\n")
run = para.add_run("5. 'One-hot encoded' text data fields.\n")
run = para.add_run("6. 'Attendance statuses reflective of a DNA or Attended appontment only (all other appointment statuses are removed).\n")

#list imputed fields
doc.add_heading("Approach to handling missing data:", 5)
doc.add_heading("Fields imputed (missing data replaced):", 5)
para = add_list_fields_kept_or_dropped_to_pydoc(doc, stage_02_Objective_0_PreProcessing_004.dna_for_report_int_float_fields_keep, 'int', 'included')
para = add_list_fields_kept_or_dropped_to_pydoc(doc, stage_02_Objective_0_PreProcessing_004.dna_for_report_str_fields_keep, 'str', 'included')
para = add_list_fields_kept_or_dropped_to_pydoc(doc, stage_02_Objective_0_PreProcessing_004.dna_for_report_datetime_fields_keep, 'date', 'included')

#list dropped fields with missing data
doc.add_heading("Fields with missing data that were removed from the analysis:", 4)
para = add_list_fields_kept_or_dropped_to_pydoc(doc, stage_02_Objective_0_PreProcessing_004.dna_for_report_int_float_fields_drop, 'int', 'excluded')
para = add_list_fields_kept_or_dropped_to_pydoc(doc, stage_02_Objective_0_PreProcessing_004.dna_for_report_str_fields_drop, 'str', 'excluded')
para = add_list_fields_kept_or_dropped_to_pydoc(doc, stage_02_Objective_0_PreProcessing_004.dna_for_report_datetime_fields_drop, 'date', 'excluded')

#sub-heading Carbon emissions section
doc.add_heading('03. Carbon Emissions from Patient Travel', 3)

#methods section
doc.add_heading("Data set description:", 4) #zxcv continue from here to build out methods text 
para = doc.add_paragraph("The final processed dataset used for the DNA Machine Learning model consists of: \n")
run = para.add_run("1. Unique appointments (not patients)\n")
run = para.add_run("2. Attended appointments only.\n")
run = para.add_run("3. 'In area' attendances only.\n")
run = para.add_run("4. Face-to-face in-person appointments only.\n")
run = para.add_run("5. Booked and walk-in appointments.\n")
run = para.add_run("6. 'Attendance statuses reflective of an Attended appontment only (all other appointment statuses are removed).\n")

#list imputed fields
doc.add_heading("Approach to handling missing data:", 5)
doc.add_heading("Fields imputed (missing data replaced):", 5)
para = add_list_fields_kept_or_dropped_to_pydoc(doc, stage_02_Objective_0_PreProcessing_004.emissions_for_report_int_float_fields_keep, 'int', 'included')
para = add_list_fields_kept_or_dropped_to_pydoc(doc, stage_02_Objective_0_PreProcessing_004.emissions_for_report_str_fields_keep, 'str', 'included')
para = add_list_fields_kept_or_dropped_to_pydoc(doc, stage_02_Objective_0_PreProcessing_004.emissions_for_report_datetime_fields_keep, 'date', 'included')

#list dropped fields with missing data
doc.add_heading("Fields with missing data that were removed from the analysis:", 4)
para = add_list_fields_kept_or_dropped_to_pydoc(doc, stage_02_Objective_0_PreProcessing_004.emissions_for_report_int_float_fields_drop, 'int', 'excluded')
para = add_list_fields_kept_or_dropped_to_pydoc(doc, stage_02_Objective_0_PreProcessing_004.emissions_for_report_str_fields_drop, 'str', 'excluded')
para = add_list_fields_kept_or_dropped_to_pydoc(doc, stage_02_Objective_0_PreProcessing_004.emissions_for_report_datetime_fields_drop, 'date', 'excluded')


'''
#check if any int fields replaced, if so, list in the doc, if not, state none replaced
if len (Objective_0_PreProcessing_dev_01_new.hea_for_report_int_float_fields_keep) > 0:
    para = doc.add_paragraph("The following lists the numeric fields the user selected to include with missing data replaced:\n")
    for field in Objective_0_PreProcessing_dev_01_new.hea_for_report_int_float_fields_keep:
        run = para.add_run(f"{field}\n")
else:
    para = doc.add_paragraph("No numeric fields with replaced missing data were included.\n")

#check if any string fields replaced, if so, list in the doc, if not, state none replaced
if len (Objective_0_PreProcessing_dev_01_new.hea_for_report_str_fields_keep) > 0:
    para = doc.add_paragraph("The following lists the text fields the user selected to include with missing data replaced:\n")
    for field in Objective_0_PreProcessing_dev_01_new.hea_for_report_str_fields_keep:
        run = para.add_run(f"{field}\n")
else:
    para = doc.add_paragraph("No text fields with replaced missing data were included.\n")

#check if any datetime fields replaced, if so, list in the doc, if not, state none replaced
if len (Objective_0_PreProcessing_dev_01_new.hea_for_report_datetime_fields_keep) > 0:
    para = doc.add_paragraph("The following lists the date time fields the user selected to include with missing data replaced:\n")    
    for field in Objective_0_PreProcessing_dev_01_new.hea_for_report_datetime_fields_keep:
        run = para.add_run(f"{field}\n")
else:
    para = doc.add_paragraph("No date time fields with replaced missing data were included.\n")

#--------------------------------------------------------------------
#now list all fields that had missing data that were removed from the analysis for each section
doc.add_heading("Fields with missing data that were removed from the analysis:", 4)

#check if any int fields replaced, if so, list in the doc, if not, state none replaced
if len (Objective_0_PreProcessing_dev_01_new.hea_for_report_int_float_fields_drop) > 0:
    para = doc.add_paragraph("The following lists the numeric fields with missing data that were excluded:\n")
    for field in Objective_0_PreProcessing_dev_01_new.hea_for_report_int_float_fields_drop:
        run = para.add_run(f"{field}\n")
else:
    para = doc.add_paragraph("No numeric fields with missing data were excluded.\n")

#check if any string fields replaced, if so, list in the doc, if not, state none replaced
if len (Objective_0_PreProcessing_dev_01_new.hea_for_report_str_fields_drop) > 0:
    para = doc.add_paragraph("The following lists the text fields with missing data that were excluded:\n")
    for field in Objective_0_PreProcessing_dev_01_new.hea_for_report_str_fields_drop:
        run = para.add_run(f"{field}\n")
else:
    para = doc.add_paragraph("No text fields with missing data were excluded.\n")

#check if any datetime fields replaced, if so, list in the doc, if not, state none replaced
if len (Objective_0_PreProcessing_dev_01_new.hea_for_report_datetime_fields_drop) > 0:
    para = doc.add_paragraph("The following lists the date time fields with missing data that were excluded:\n")    
    for field in Objective_0_PreProcessing_dev_01_new.hea_for_report_datetime_fields_drop:
        run = para.add_run(f"{field}\n")
else:
    para = doc.add_paragraph("No date time fields with missing data were excluded.\n")
'''

#-----------------------------------------------------------
#Appendix 2 - all HEA charts
add_appendix(doc, 'HEA Charts', 1, 'This appendix includes charts for each condition in scope, identifying any significant variation in expressed need when compared to the population profile.', preferred_font_name, preferred_font_size)

doc.add_paragraph(' The charts flag statistically significant findings with the word "SIG" at the end of the bar. Where the difference is not statistically significant, this is indicated by the term "ns".')

#loop through each demographic feature, and each chart for each condition in scope:
#Age
doc.add_heading('Age:', 4)

for condition in list_of_conditions:
    doc.add_heading(f'{condition}:', 5)
    inline_shape, inline_shape.width, inline_shape.height = change_image_size_pydoc(doc, 0.5, dict_condition_Age_chart_file_paths[condition])

#Ethnicity
doc.add_heading('Ethnicity:', 4)
for condition in list_of_conditions:
    doc.add_heading(f'{condition}:', 5)
    inline_shape, inline_shape.width, inline_shape.height = change_image_size_pydoc(doc, 0.5, dict_condition_Ethnicity_chart_file_paths[condition])

#IMD
doc.add_heading('IMD:', 4)
for condition in list_of_conditions:
    doc.add_heading(f'{condition}:', 5)
    inline_shape, inline_shape.width, inline_shape.height = change_image_size_pydoc(doc, 0.5, dict_condition_IMD_chart_file_paths[condition])

doc.add_heading('Gender:', 4)

#use condition logic to check whether the user has modelled all conditions for all persons. Doing so avoids an error when trying to subsequently read in the image file paths for inclusion in the report. 
user_stated_gender_for_all_conditions_was_all_persons = True
subset_list_conditions = []
for condition in list_of_conditions:
    if dict_pop_gender_for_each_condition_text[condition] == 'Persons':
        subset_list_conditions.append(condition)
    else:
        user_stated_gender_for_all_conditions_was_all_persons = False

if user_stated_gender_for_all_conditions_was_all_persons != True:
    doc.add_paragraph('In undertaking the HEA analysis, the user of the decision support tool indicated not all conditions modelled for see all genders. For this reason, only those conditions modelled that have been indicated as seeing all genders are included here.')

for condition in subset_list_conditions:
    doc.add_heading(f'{condition}:', 5)
    inline_shape, inline_shape.width, inline_shape.height = change_image_size_pydoc(doc, 0.5, dict_condition_Gender_chart_file_paths[condition])

#Need to add gender once solved how to resize images where not all conditions may be present. likely needs a bespoke function ? 

#THIS ADD TABLE FUNCTION ISNT WORKING - STYLE NAME NOT RECOGNISED
#add_df_as_table(preferred_font_name, doc, report_table, dict_list_table_styles['light_grey_01'])
#add_df_as_table(preferred_font_name, 'zExample_Report.docx', report_table, dict_list_table_styles['grey_01'])

#List Table 6 Colorful Accent 1
#new_table = doc.add_table(report_table.shape[0]+1, report_table.shape[1], style=dict_list_table_styles['yellow_06'])

"""
report_table = dict_condition_df_sig_findings['GUM']

new_table = doc.add_table(report_table.shape[0]+1, report_table.shape[1], style=dict_list_table_styles['yellow_06'])

# add the header rows.
for j in range(report_table.shape[-1]):
    new_table.cell(0,j).text = report_table.columns[j]

# add the rest of the data frame
for i in range(report_table.shape[0]):
    for j in range(report_table.shape[-1]):
        new_table.cell(i+1,j).text = str(report_table.values[i,j])

for row in new_table.rows:
    for cell in row.cells:
        paragraphs = cell.paragraphs
        for paragraph in paragraphs:
            for run in paragraph.runs:
                font = run.font
                font.name = preferred_font_name
                font.size= Pt(8)
"""


#-----------------------------------------
'''
#Appendix 2 - summary of HEA method and assumptions for unmet need
add_appendix(doc, "Summary of the HEA method used and assumptions made", 1, 'This appendix includes a summary overview of the method and assumptions underpinning the HEA and modelled unmet need summarised in the main body of the report.', preferred_font_name, preferred_font_size)
paraObj2 = doc.add_heading("Method:", 2)
paraObj3 = doc.add_heading("Assumptions:", 2)
'''

#-----------------------------------------
#Appendix 3 - summary of all ML models performance
add_appendix(doc, "Summary of Machine Learning model variants' performance", 1, 'This appendix includes a summary table consisting of the metrics for each of the 9 machine learning model variants summarising how they each performed when trying to predict whether a patient would DNA or not.', preferred_font_name, preferred_font_size)
#add_df_as_table_test(preferred_font_name, doc, dict_dna_higher_lower_chance[direction], chosen_style)

#source for ROC AUC interpretation:
# https://www.tc.columbia.edu/elda/blog/content/receiver-operating-characteristic-roc-area-under-the-curve-auc/
paraObj4 = doc.add_paragraph()
run = paraObj4.add_run('To interpret how well a classifcation model is performing, the ROC AUC measure is often used, with the following as a guideline (D Agostino, Rodgers, & Mauck, 2018):\n')
doc.add_paragraph('AUC above 0.85 means high classification accuracy')
doc.add_paragraph('AUC between 0.75 and 0.85 moderate accuracy')
doc.add_paragraph('AUC less than 0.75 low accuracy\n')

df_of_metrics = pd.read_csv(f"{cwd}/Assets_produced_by_code/03_DNA_ML_assets/df005_AllMetricsCombinedTransposed.csv")
df_of_metrics.rename(columns={'Metric': 'Model variant', 'predicted_positive_rate': 'predicted pos. rate', 'observed_positive_rate': 'observed pos. rate'}, inplace=True)
add_df_as_table(preferred_font_name, doc, df_of_metrics, chosen_style)


#-----------------------------------------
# appendix 4: Explanation of Terms Used / glossary 

#add a heading. Heading "level" can be determined by the interger param after the string
#0 = title, 1 = top header, 4 = smallest subheading
add_appendix(doc, 'Glossary', 1, 'This appendix includes a brief explanation of frequently used terms throughout this report.', preferred_font_name, preferred_font_size)

#list terms and give explanation
# DNA
paraObj2 = doc.add_paragraph()
run = paraObj2.add_run("Did Not Attend (DNA): ").bold=True
run = paraObj2.add_run("A commonly used label use to describe an appointment accepted by a patient, but not attended.")

# Expressed Need
paraObj2 = doc.add_paragraph()
run = paraObj2.add_run("Expressed Need: ").bold=True
run = paraObj2.add_run("Felt need turned into action, e.g., help seeking behaviour such as actively going to the doctor to have the pain in their arm checked out.")

# Felt Need
paraObj2 = doc.add_paragraph()
run = paraObj2.add_run("Felt Need: ").bold=True
run = paraObj2.add_run("Need perceived by an individual, e.g., a person who feels they have a pain in their arm.")

# HEA
paraObj2 = doc.add_paragraph()
run = paraObj2.add_run("'Health Equity Assessment (HEA)': ").bold=True
run = paraObj2.add_run("This is a public health tool used to identify inequalities present in service uptake and / or outcomes. The outputs from an HEA should provide the area of focus for any subsequent work to improve inequalities.")

# Logistic Regression
paraObj2 = doc.add_paragraph()
run = paraObj2.add_run("Logistic Regression: ").bold=True
run = paraObj2.add_run("A method used to fit a regression model with a binary response variable, e.g., 'attendance status' with the outcomes 'attended appointment' and 'DNA appointment'. To determine how well a logistic regression 'fits' a dataset two metrics are commonly used: sensitivity and specificity.")

# Machine Learning
paraObj2 = doc.add_paragraph()
run = paraObj2.add_run("Machine Learning (ML): ").bold=True
run = paraObj2.add_run("A branch of Artificial Intelligence, ML is a field that seeks to understand and build methods that 'learn'. This commonly involves building a model using sample (or, 'training') data, so that the mode can make predictions about unseen 'testing' data, without being explicity programmed to do so (source: Wikipedia). ")

# One-hot encoding
paraObj2 = doc.add_paragraph()
run = paraObj2.add_run("One-hot encoding: ").bold=True
run = paraObj2.add_run("A method of converting text fields in a data set into columns consisting of 0's and 1's. For example, imagine a 'disabled' field, consisting of 2 possible values: 'yes' where the patient for that data row was disabled, and 'no' where they weren't. When one-hot-encoded, would be converted to 2 columns, with 1 column representing 'disabled-yes' and the other representing 'disabled-no', with the value of 1 present when the status for that row is True, and 0 when False.")

# ROC Curve
paraObj2 = doc.add_paragraph()
run = paraObj2.add_run("Receiver Characteristic Curve (ROC curve): ").bold=True
run = paraObj2.add_run("A plot displaying the sensitivity and specificity of a logistic regression model, commonly used to visualise these two metrics. Interpreted using the 'Area Under the Curve of the ROC' (or, AUC-ROC), where an AUC-ROC of 1.0 represents a perfect model (so, the aim is to get AUC-ROC values as close to 1 as possible). If a model has an AUC-ROC of 0.5 this can be interpreted as effectively being no better at predicting the outcome than flipping a coin.")

# Sensitivity
paraObj2 = doc.add_paragraph()
run = paraObj2.add_run("Sensitivity: ").bold=True
run = paraObj2.add_run("The true positive rate i.e., the probability that model predicts a positive outcome, when the outcome truly is positive.")

# Specificity
paraObj2 = doc.add_paragraph()
run = paraObj2.add_run("Specificity: ").bold=True
run = paraObj2.add_run("The true negative rate i.e., the probability that model predicts a negative outcome, when the outcome truly is negative.")


#-----------------------------------------
#Appendix 5 - Acknowledgements and references
add_appendix(doc, "Acknowledgements and references", 1, '', preferred_font_name, preferred_font_size)
para = doc.add_heading("Ackknowledgements:", 2)
para = doc.add_paragraph("The following acknowledges the helpful advice and input to support the production of the automated tool that produced this report. The authors respectfully ask that you familiarise yourself with the attribution and use guidance (URL's below) and ensure appropriate attribution and acknowledgement is made if you share any of the outputs from this document.")
para = doc.add_paragraph('1. All of the HSMA Team at PenARC, in particular Dan Chalk, and Mike Allen whose titanic work has been adapted here under the MIT licence (https://github.com/MichaelAllen1966/titanic)')
para = doc.add_paragraph('2. Derbyshire Community Health Services NHS FT, The University of Exeter, and West Sussex County Council.')
para = doc.add_paragraph('3. The CodeWith online community, in particular Drew Morgan.')
para = doc.add_paragraph('4. The Open Route Service https://openrouteservice.org/ (© openrouteservice.org by HeiGIT)' )
para = doc.add_paragraph('5. Open Street Map https://www.openstreetmap.org/copyright (Map data © OpenStreetMap contributors)')
para = doc.add_paragraph('6. The Office for National Statistics')
para = doc.add_paragraph('7. Nomis (official census and labour market statistics) https://www.nomisweb.co.uk/ ')
para = doc.add_paragraph('8. The Open Geography Portal https://geoportal.statistics.gov.uk/ ')
para = doc.add_paragraph('9. The Ministry of Housing, Communities and Local Government https://geoportal.statistics.gov.uk/ ')
para = doc.add_paragraph('')
para = doc.add_paragraph('For further information on the techniques used to produce this document, please refer to https://github.com/hsma4')

paraObj3 = doc.add_heading("References:", 2)
paraObj3 = doc.add_paragraph("ROC AUC classification cut-off points https://www.tc.columbia.edu/elda/blog/content/receiver-operating-characteristic-roc-area-under-the-curve-auc/")

# ------------------------------

report_production_time = round(((time.time() - report_start_time)/60),2)



print("\n--------------------------------------------------")
print("\nSummary of run time for each stage:")
print(f"Pre-Processing code run successfully in {processing_run_time} minutes.")
if ML_model_run == True:
    print(f"ML code run successfully in {ml_run_time} minutes.")
print(f"HEA run successfully in {hea_run_time} minutes.")
print(f"Carbon emissions model run successfully in {carbon_run_time} minutes.")
print(f"Maps produced in {maps_run_time} minutes.")
print(f'Report produced in {report_production_time} minutes and saved successfully\n')
print("--------------------------------------------------\n")

#save document object to a file
doc.save(f'z{service_name}_Automated_Report.docx')


#-------------------------------------------------------
#-------------------------------------------------------
#<<<<<     Produce report code ends     >>>>>
#-------------------------------------------------------






print("\nEnd of model. Please refer to the output report containing all outputs based on the parameters you have used. Thank you.")